from openpyxl import load_workbook
from datetime import datetime
from docx import Document
from fpdf import FPDF

planilha_fornecedores = load_workbook('./fornecedores.xlsx')
pagina_fornecedores = planilha_fornecedores['Sheet1']

for linha in pagina_fornecedores.iter_rows(min_row=2, values_only=True):
    nome_empresa, endereco, cidade, estado, cep, telefone, email, setor = linha

    arquivo_word = Document()
    arquivo_word.add_heading('Contrato de prestação de serviço', 0)

    texto_contrato = f"""

    contrato de prestação de serviços é feito entre {nome_empresa}, com endereço em {endereco},
    {cidade},  CEP Nº {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

    Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

    1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

    2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

    3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

    4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

    Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

    FORNECEDOR: {nome_empresa}
    E-mail: {email}

    CONTRATANTE: LUCAS LTDA
    E-mail: lucasltda@hotmail.com

    Bahia, {datetime.now().strftime('%d/%m/%Y, às %H:%M:%S')}
"""
    arquivo_word.add_paragraph(texto_contrato)
    arquivo_word.save(f'./contratos_{nome_empresa}.docx')

    doc = Document(f'./contratos_{nome_empresa}.docx')
    texto = ""

    for paragraph in doc.paragraphs:
        texto += paragraph.text + "\n"

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial",size=12)
    pdf.multi_cell(0,10,texto)
    pdf.output(f'./contratos_{nome_empresa}.pdf')

